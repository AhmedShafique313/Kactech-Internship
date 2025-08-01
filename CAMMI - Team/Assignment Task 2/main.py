from docx import Document

def fix_heading_paragraphs(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    i = 0
    paragraphs = doc.paragraphs
    while i < len(paragraphs):
        current_text = paragraphs[i].text.strip()

        # Check if current line ends with ":" and next line is a continuation
        if current_text.endswith(":") and i + 1 < len(paragraphs):
            next_text = paragraphs[i + 1].text.strip()

            if next_text:  # Ensure the next paragraph is not empty
                # Merge heading and paragraph
                combined_text = f"{current_text} {next_text}"
                new_doc.add_paragraph(combined_text)
                i += 2  # Skip the next line as it's already processed
                continue

        # Otherwise, keep the paragraph as it is
        new_doc.add_paragraph(current_text)
        i += 1

    new_doc.save(output_path)
    print(f"Formatted document saved to {output_path}")

# Example usage:
fix_heading_paragraphs("input.docx", "output.docx")