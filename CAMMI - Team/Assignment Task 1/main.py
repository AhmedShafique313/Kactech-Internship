import os, re
from docx import Document
from groq import Groq
from dotenv import load_dotenv
load_dotenv(dotenv_path=".env")

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def generate_blog(prompt):
    messages = [
        {"role": "system", "content": "You are a professional blog writer."},
        {"role": "user", "content": f"Write a well-structured, properly formatted, logically structured and engaging blog article based on the following input:\n\n{prompt}"}
    ]
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=messages,
        temperature=0.7,
        max_tokens=800
    )
    output_msg = response.choices[0].message.content.strip()
    return output_msg

def read_input_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
    
def save_to_docx(title, content, output_path):
    doc = Document()
    doc.add_heading(title, level=1)

    paragraphs = content.split('\n')
    for para in paragraphs:
        line = para.strip()
        if not line:
            continue

        # Check for markdown-style bold heading
        bold_heading_match = re.match(r'^\*\*(.+?)\*\*$', line)
        is_numbered_list = re.match(r'^\d+\.\s', line)

        is_heading = (
            bold_heading_match or
            (len(line.split()) < 12 and (
                line.endswith(':') or
                line.istitle() or
                line.isupper() or
                line.lower().startswith("conclusion")
            ))
        )

        if bold_heading_match:
            clean_line = bold_heading_match.group(1).strip()
            doc.add_heading(clean_line, level=2)

        elif is_heading:
            clean_line = line.strip('#:* ').strip('*')
            doc.add_heading(clean_line, level=2)

        elif is_numbered_list:
            # Add numbered list items
            doc.add_paragraph(line, style='List Number')

        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

    
if __name__ == "__main__":
    input_path = r"C:\Users\HA LYSM\Documents\Projects\Auto Blog Generator\input_file.txt"
    output_path = r"C:\Users\HA LYSM\Documents\Projects\Auto Blog Generator\generated_blog.docx"
    input_text = read_input_text(input_path)
    blog_content = generate_blog(input_text)
    blog_title = "The Impact of AI on Modern Society"
    save_to_docx(blog_title, blog_content, output_path)