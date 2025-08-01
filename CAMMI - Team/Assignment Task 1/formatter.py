import re
from docx import Document

def structure_and_save_docx(input_path, output_path):
    original_doc = Document(input_path)
    new_doc = Document()
    paragraphs = [p.text.strip() for p in original_doc.paragraphs if p.text.strip()]

    new_doc.add_heading(paragraphs[0], level=1)

    for line in paragraphs[1:]:
        # --- Handle headings ---
        if line.lower() in ["introduction", "conclusion"] or (line.istitle() and len(line.split()) < 10):
            new_doc.add_heading(line, level=2)
            continue

        # --- Handle subheading like "- **Healthcare:** something"
        bold_heading = re.match(r'^-\s*\*\*(.+?):\*\*\s*(.*)', line)
        if bold_heading:
            heading = bold_heading.group(1).strip()
            paragraph = bold_heading.group(2).strip()
            new_doc.add_heading(heading, level=3)
            if paragraph:
                new_doc.add_paragraph(paragraph)
            continue

        # --- Bullet point ---
        if line.startswith("- "):
            new_doc.add_paragraph(line[2:].strip(), style='List Bullet')
            continue

        # --- Numbered list ---
        if re.match(r'^\d+\.\s+', line):
            new_doc.add_paragraph(line.strip(), style='List Number')
            continue

        # --- Default paragraph ---
        new_doc.add_paragraph(line)

    new_doc.save(output_path)
    print(f"Structured document saved to: {output_path}")

if __name__ == "__main__":
    input_path = r"C:\Users\HA LYSM\Documents\Projects\Auto Blog Generator\generated_blog.docx"
    output_path = r"C:\Users\HA LYSM\Documents\Projects\Auto Blog Generator\structured_blog.docx"

    structure_and_save_docx(input_path, output_path)
